"""
Script generation service.
Takes generated course content and produces a presenter/narrator Word document script.
"""

import io
import logging
from typing import Any, Dict, Optional
import httpx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import settings

logger = logging.getLogger(__name__)

CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"
CLAUDE_MODEL = "claude-sonnet-4-6"


class ScriptService:
    """Generates presenter scripts from course content."""

    def __init__(self):
        self.api_key = settings.CLAUDE_API_KEY

    async def generate_script(self, course_content: Dict[str, Any]) -> bytes:
        """
        Generate a presenter script for the course and return as .docx bytes.

        Args:
            course_content: The course content dict (stored in Course.content)

        Returns:
            .docx file as bytes
        """
        course_title = course_content.get("title", "Untitled Course")
        logger.info(f"Generating presenter script for: {course_title}")

        script_text = await self._call_claude_for_script(course_content)
        docx_bytes = self._build_docx(course_title, script_text)

        logger.info(f"Script generated for: {course_title}")
        return docx_bytes

    async def _call_claude_for_script(self, course_content: Dict[str, Any]) -> str:
        """Call Claude to write the presenter script."""
        prompt = self._build_script_prompt(course_content)

        headers = {
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
            "x-api-key": self.api_key,
        }

        payload = {
            "model": CLAUDE_MODEL,
            "max_tokens": 8192,
            "messages": [{"role": "user", "content": prompt}],
        }

        async with httpx.AsyncClient() as client:
            response = await client.post(
                CLAUDE_API_URL,
                json=payload,
                headers=headers,
                timeout=180.0,
            )

        if response.status_code != 200:
            logger.error(f"Claude API error: {response.status_code} - {response.text}")
            raise Exception(f"Claude API error: {response.status_code}")

        return response.json().get("content", [{}])[0].get("text", "")

    def _build_script_prompt(self, course_content: Dict[str, Any]) -> str:
        """Build the prompt for script generation."""
        title = course_content.get("title", "Untitled Course")
        description = course_content.get("description", "")
        modules = course_content.get("modules", [])

        modules_summary = []
        for mod in modules:
            lessons = mod.get("lessons", [])
            lesson_titles = [l.get("title", "") for l in lessons]
            modules_summary.append(
                f"Module {mod.get('id', '')}: {mod.get('title', '')} — "
                f"Lessons: {', '.join(lesson_titles)}"
            )

        course_json_str = str(course_content)[:6000]  # Limit to avoid token overflow

        return f"""You are writing a presenter/narrator script for a training course.

Course Title: {title}
Course Description: {description}

Modules:
{chr(10).join(modules_summary)}

Full course content (JSON):
{course_json_str}

Write a complete, professional presenter script for this course. Requirements:
- Write in flowing, spoken-word style — how a confident presenter would actually say it
- NOT bullet points — full sentences and natural speech patterns
- Include a warm introduction to the whole course
- For each module: a transition/introduction sentence setting up the module
- For each lesson: a complete spoken narration covering the key content and learning outcomes
- Include natural transitions between lessons ("Now that we've covered X, let's move on to...")
- Close with a summary and call to action

Format your response with clear section markers:
[INTRO]
(course introduction script)

[MODULE 1: Module Title]
[LESSON 1.1: Lesson Title]
(lesson script)

[LESSON 1.2: Lesson Title]
(lesson script)

[MODULE 2: Module Title]
...and so on

[CLOSE]
(closing script)

Write the full script now:"""

    def _build_docx(
        self,
        course_title: str,
        script_text: str,
    ) -> bytes:
        """Build the .docx file from the script text."""
        doc = Document()

        # Title page
        title_para = doc.add_heading(course_title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph("Presenter Script")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        doc.add_page_break()

        # Parse and render the script sections
        # Handles both [SECTION] markers and markdown headings (#, ##, ###)
        lines = script_text.split("\n")
        current_para = None

        for line in lines:
            line = line.strip()
            if not line:
                current_para = None
                continue

            # Skip separator lines
            if set(line) <= set("-*_ ") and len(line) >= 2:
                current_para = None
                continue

            # Markdown headings
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                current_para = None
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
                current_para = None
            elif line.startswith("# "):
                # Top-level markdown heading — skip if it duplicates the title
                label = line[2:].strip()
                if label.lower() not in (course_title.lower(), f"{course_title.lower()} — full presenter script", "presenter script"):
                    doc.add_heading(label, level=1)
                current_para = None
            # Legacy [SECTION] markers
            elif line.startswith("[MODULE"):
                doc.add_heading(line.strip("[]"), level=1)
                current_para = None
            elif line.startswith("[LESSON"):
                doc.add_heading(line.strip("[]"), level=2)
                current_para = None
            elif line.startswith("[INTRO]"):
                doc.add_heading("Introduction", level=1)
                current_para = None
            elif line.startswith("[CLOSE]"):
                doc.add_heading("Closing", level=1)
                current_para = None
            else:
                # Normal script content — strip leading markdown bold markers
                text = line.lstrip("*").rstrip("*").strip()
                if not text:
                    continue
                if current_para is None:
                    current_para = doc.add_paragraph(text)
                    current_para.style = doc.styles["Normal"]
                else:
                    current_para.add_run(" " + text)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
